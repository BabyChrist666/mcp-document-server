"""
Document parsers for PDF, DOCX, and plaintext files.
"""

import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Optional, Dict, Any, List
from pathlib import Path


@dataclass
class DocumentMetadata:
    """Extracted document metadata."""
    title: str = ""
    author: str = ""
    pages: int = 0
    word_count: int = 0
    file_type: str = ""
    file_size_bytes: int = 0
    file_path: str = ""
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    text: str
    metadata: DocumentMetadata
    pages: List[str] = field(default_factory=list)


class BaseParser(ABC):
    """Abstract base parser for documents."""

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document and return extracted text and metadata."""
        pass

    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """Check if this parser supports the given file type."""
        pass

    def _base_metadata(self, file_path: str, file_type: str) -> DocumentMetadata:
        """Create base metadata from file path."""
        path = Path(file_path)
        return DocumentMetadata(
            title=path.stem,
            file_type=file_type,
            file_size_bytes=path.stat().st_size if path.exists() else 0,
            file_path=str(path.absolute()),
        )


class PDFParser(BaseParser):
    """Parse PDF documents using PyPDF2."""

    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".pdf")

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing: pip install PyPDF2")

        reader = PdfReader(file_path)
        metadata = self._base_metadata(file_path, "pdf")

        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        full_text = "\n\n".join(pages)
        metadata.pages = len(reader.pages)
        metadata.word_count = len(full_text.split())

        # Extract PDF metadata
        pdf_info = reader.metadata
        if pdf_info:
            metadata.title = pdf_info.get("/Title", metadata.title) or metadata.title
            metadata.author = pdf_info.get("/Author", "") or ""

        return ParsedDocument(text=full_text, metadata=metadata, pages=pages)


class DocxParser(BaseParser):
    """Parse DOCX documents using python-docx."""

    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing: pip install python-docx")

        doc = Document(file_path)
        metadata = self._base_metadata(file_path, "docx")

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        metadata.word_count = len(full_text.split())

        core_props = doc.core_properties
        if core_props:
            metadata.title = core_props.title or metadata.title
            metadata.author = core_props.author or ""

        return ParsedDocument(text=full_text, metadata=metadata)


class TextParser(BaseParser):
    """Parse plaintext files."""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml"}

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> ParsedDocument:
        metadata = self._base_metadata(file_path, Path(file_path).suffix.lstrip("."))

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        metadata.word_count = len(text.split())
        return ParsedDocument(text=text, metadata=metadata)


class DocumentParserRegistry:
    """Registry of all available document parsers."""

    def __init__(self):
        self._parsers: List[BaseParser] = [
            PDFParser(),
            DocxParser(),
            TextParser(),
        ]

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document using the appropriate parser."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        for parser in self._parsers:
            if parser.supports(file_path):
                return parser.parse(file_path)

        raise ValueError(
            f"Unsupported file type: {Path(file_path).suffix}. "
            f"Supported: PDF, DOCX, TXT, MD, CSV, LOG, JSON, YAML"
        )

    def supported_types(self) -> List[str]:
        return ["pdf", "docx", "txt", "md", "csv", "log", "json", "yaml", "yml"]
